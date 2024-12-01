from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, RGBColor, Inches
import random
import json

import generic_warm_and_fuzzies

from emojis import EMOJIS

WARM_AND_FUZZIES = {}

# Optionally, generate generic warm and fuzzies for everyone in the contacts list.
# This is useful if you want to send a warm and fuzzy to everyone, even if they didn't
# recieve one. 
# To modify the generic message, open `generic_warm_and_fuzzies.py` and change the message.
# If you do not want to generate generic warm and fuzzies, comment out the line below.
WARM_AND_FUZZIES = generic_warm_and_fuzzies.generate(contacts_filename="contacts.xlsx")

# INDEXES i.e. what column is this data in your spreadsheet
TIMESTAMP = 0
MESSAGE = 16

# Replace 'responses.xlsx' with your own file, or rename your own file
workbook = load_workbook(filename="responses.xlsx")
sheet = workbook.active

# Go through each row until hit an empty row
# This finds the name of the recipient, may need to change min_row, max_col
# and the name_start, name_end
name_start = 3
name_end = 16
for row in sheet.iter_rows(min_row=2, max_col=18, max_row=sheet.max_row):
    for col in range(name_start, name_end):
        if row[col].value != None:
            name = row[col].value.strip().upper()
            
    print(name)
    if name not in WARM_AND_FUZZIES:
        WARM_AND_FUZZIES[name] = []

    WARM_AND_FUZZIES[name].append(row[MESSAGE].value)

# print(json.dumps(WARM_AND_FUZZIES, indent=2))

# print(len(WARM_AND_FUZZIES), "recipients...")

# Dump each person into a LaTeX file
for name in WARM_AND_FUZZIES:
    doc = Document()

    doc.styles["Heading 1"].font.name = "CMU Serif"
    doc.styles["Heading 1"].font.size = Pt(24)
    doc.styles["Heading 1"].font.color.rgb = RGBColor(0, 0, 0)

    doc.styles["Normal"].font.name = "CMU Serif"
    doc.styles["Normal"].paragraph_format.space_before = Pt(12)

    p = doc.add_paragraph()
    p.alignment = 1
    run = p.add_run("")
    run.add_picture("slacklogo.png", width=Inches(1.5))

    doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = 1
    run = t.add_run(name)
    run.font.size = Pt(24)
    run.bold = True

    h = doc.add_paragraph()
    h.alignment = 1
    h.add_run("UNSW CSESOC 2024").italic = True

    for message in WARM_AND_FUZZIES[name]:
        divider = doc.add_paragraph(random.choice(EMOJIS))
        divider.alignment = 1
        doc.add_paragraph(message.strip().strip("\n"))

    file_name =f"output_warm_and_fuzzies/{name}.docx"

    # save each file as a .docx
    doc.save(file_name)
